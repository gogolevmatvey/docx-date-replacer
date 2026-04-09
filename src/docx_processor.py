"""
Модуль для обработки документов .docx.
"""

import os
import logging
from typing import List, Tuple
from copy import deepcopy
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from lxml.etree import _Element

from .config import DEFAULT_CONFIG
from .date_replacer import DateReplacer

logger = logging.getLogger(__name__)

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _extract_cell_text(tc) -> str:
    """Извлечение полного текста из ячейки таблицы."""
    return ''.join(t.text for t in tc.iter(qn('w:t')) if t.text)


def _copy_formatting_from_run(r_element: _Element) -> _Element:
    """Копирование форматирования (шрифт + размер) из исходного XML-элемента w:r."""
    rPr_new = OxmlElement('w:rPr')

    rPr = r_element.find(qn('w:rPr'))
    if rPr is not None:
        # Копируем шрифт
        orig_fonts = rPr.find(qn('w:rFonts'))
        if orig_fonts is not None:
            rPr_new.append(deepcopy(orig_fonts))

        # Копируем размер шрифта из оригинала
        orig_sz = rPr.find(qn('w:sz'))
        if orig_sz is not None:
            rPr_new.append(deepcopy(orig_sz))

    # Если шрифт не найден — ставим Times New Roman
    if rPr_new.find(qn('w:rFonts')) is None:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(f'{NS}ascii', 'Times New Roman')
        rFonts.set(f'{NS}hAnsi', 'Times New Roman')
        rFonts.set(f'{NS}cs', 'Times New Roman')
        rPr_new.append(rFonts)

    return rPr_new


class DocxProcessor:
    """Класс для обработки документов .docx."""

    def __init__(self, date_replacer: DateReplacer):
        """
        Инициализация процессора документов.
        """
        self.date_replacer = date_replacer

    def load_document(self, filepath: str) -> DocumentType:
        """Загрузка документа .docx."""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Файл не найден: {filepath}")
        return Document(filepath)

    def save_document(self, doc: DocumentType, filepath: str) -> None:
        """Сохранение документа .docx."""
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        logger.info(f"Документ сохранен: {filepath}")

    def has_date_in_first_paragraphs(self, doc: DocumentType) -> bool:
        """Проверка наличия даты на первой странице."""
        max_paragraphs = DEFAULT_CONFIG.first_page_paragraphs
        max_cells = 100

        logger.debug(f"Поиск даты в первых {max_paragraphs} параграфах...")

        # Параграфы
        for paragraph in doc.paragraphs[:max_paragraphs]:
            if paragraph.text.strip() and self.date_replacer.find_date(paragraph.text):
                return True

        # Таблицы — поиск через XML для максимального охвата
        cell_count = 0
        for tc in doc._element.body.iter(qn('w:tc')):
            if cell_count >= max_cells:
                break
            cell_text = _extract_cell_text(tc)
            if cell_text.strip():
                if self.date_replacer.find_date(cell_text):
                    return True
                cell_count += 1

        logger.debug("Дата не найдена на первой странице.")
        return False

    def process_paragraphs(self, doc: DocumentType, first_page_only: bool = True) -> Tuple[int, int]:
        """Обработка параграфов документа."""
        processed = 0
        replaced = 0

        max_paragraphs = DEFAULT_CONFIG.first_page_paragraphs if first_page_only else len(doc.paragraphs)
        paragraphs_to_process = doc.paragraphs[:max_paragraphs]

        for paragraph in paragraphs_to_process:
            if not paragraph.text.strip():
                continue

            if not self.date_replacer.find_date(paragraph.text):
                processed += 1
                continue

            count = self._process_paragraph_runs(paragraph)
            replaced += count
            processed += 1

            if first_page_only and replaced > 0:
                break

        return (processed, replaced)

    def _process_paragraph_runs(self, paragraph) -> int:
        """Замена даты в run-ах одного параграфа. Возвращает количество замен."""
        if not paragraph.text.strip():
            return 0
        if not self.date_replacer.find_date(paragraph.text):
            return 0

        # 1. Пытаемся заменить в каждом run по отдельности
        for run in paragraph.runs:
            new_text, is_changed = self.date_replacer.replace_date(run.text)
            if is_changed:
                run.text = new_text
                return 1

        # 2. Дата разбита по run-ам — сложная замена через XML
        full_text = paragraph.text
        match = self.date_replacer.search_pattern().search(full_text)
        if not match:
            return 0

        start_idx = match.start()
        end_idx = match.end()
        new_date_text = self.date_replacer.new_date

        current_pos = 0
        runs_to_modify = []

        for run in paragraph.runs:
            run_len = len(run.text)
            run_start = current_pos
            run_end = current_pos + run_len

            if run_start < end_idx and run_end > start_idx:
                runs_to_modify.append((run, run_start, run_end))

            current_pos = run_end

        if not runs_to_modify:
            return 0

        first_r = runs_to_modify[0][0]._element
        new_r = OxmlElement('w:r')
        rPr_new = _copy_formatting_from_run(first_r)
        new_r.append(rPr_new)

        new_t = OxmlElement('w:t')
        new_t.text = new_date_text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)

        first_r.addprevious(new_r)

        for run, _, _ in runs_to_modify:
            run.text = ''
            parent = run._element.getparent()
            has_text = any(
                child.text and child.text.strip()
                for child in run._element.iter(qn('w:t'))
            )
            if not has_text and parent is not None:
                parent.remove(run._element)

        return 1

    def process_tables(self, doc: DocumentType, first_page_only: bool = True) -> Tuple[int, int]:
        """Обработка таблиц через XML-параграфы ячеек."""
        processed = 0
        replaced = 0
        max_cells = DEFAULT_CONFIG.first_page_paragraphs if first_page_only else 1000

        cell_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if processed >= max_cells:
                        return (processed, replaced)

                    # Обрабатываем каждый XML-параграф (w:p) внутри ячейки
                    for p_xml in cell._element.iter(qn('w:p')):
                        cell_count += 1
                        if cell_count > max_cells:
                            return (processed, replaced)

                        # Собираем текст параграфа
                        para_text = ''.join(
                            t.text for t in p_xml.iter(qn('w:t')) if t.text
                        )
                        if not para_text.strip():
                            continue

                        if not self.date_replacer.find_date(para_text):
                            continue

                        # Создаём временный объект-обёртку для _process_paragraph_runs
                        # Т.к. _process_paragraph_runs работает с paragraph.runs,
                        # а здесь у нас чистый XML, делаем замену напрямую
                        count = self._replace_in_xml_paragraph(p_xml)
                        if count > 0:
                            replaced += count
                            processed += 1

        return (processed, replaced)

    def _replace_in_xml_paragraph(self, p_xml) -> int:
        """Замена даты в XML-элементе w:p. Возвращает количество замен."""
        # Собираем текст и элементы
        text_elements = list(p_xml.iter(qn('w:t')))
        if not text_elements:
            return 0

        full_text = ''.join(t.text for t in text_elements if t.text)
        match = self.date_replacer.search_pattern().search(full_text)
        if not match:
            return 0

        new_date_text = self.date_replacer.new_date
        start_idx = match.start()
        end_idx = match.end()

        # Находим элементы, попадающие в диапазон даты
        current_pos = 0
        elems_to_modify = []

        for t_elem in text_elements:
            t_len = len(t_elem.text or '')
            t_start = current_pos
            t_end = current_pos + t_len

            if t_start < end_idx and t_end > start_idx:
                elems_to_modify.append((t_elem, t_start, t_end))

            current_pos = t_end

        if not elems_to_modify:
            return 0

        # Копируем форматирование из первого затронутого элемента
        first_r = elems_to_modify[0][0].getparent()
        new_r = OxmlElement('w:r')
        rPr_new = _copy_formatting_from_run(first_r)
        new_r.append(rPr_new)

        new_t = OxmlElement('w:t')
        new_t.text = new_date_text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)

        first_r.addprevious(new_r)

        # Очищаем старые элементы
        for t_elem, _, _ in elems_to_modify:
            t_elem.text = ''
            r = t_elem.getparent()
            parent = r.getparent()
            if parent is not None:
                has_text = any(
                    child.text and child.text.strip()
                    for child in r.iter(qn('w:t'))
                )
                if not has_text:
                    parent.remove(r)

        return 1

    def process_document(self, input_path: str, output_path: str) -> Tuple[bool, str, int]:
        """Полная обработка документа."""
        try:
            logger.info(f"Обработка файла: {input_path}")
            doc = self.load_document(input_path)

            has_old_date = self.has_date_in_first_paragraphs(doc)

            if not has_old_date:
                self.save_document(doc, output_path)
                return (True, "Файл скопирован (дата не найдена)", 0)

            _, para_replaced = self.process_paragraphs(doc, first_page_only=True)
            _, table_replaced = self.process_tables(doc, first_page_only=True)

            total_replaced = para_replaced + table_replaced

            self.save_document(doc, output_path)

            if total_replaced == 0:
                return (True, "Файл скопирован (замен не потребовалось)", 0)

            return (True, f"Заменено дат: {total_replaced}", total_replaced)

        except Exception as e:
            logger.error(f"Ошибка при обработке {input_path}: {str(e)}")
            return (False, f"Ошибка: {str(e)}", 0)

    def copy_folder_structure(self, src_root: str, dst_root: str) -> None:
        """Копирование структуры папок."""
        for dirpath, _, _ in os.walk(src_root):
            rel_path = os.path.relpath(dirpath, src_root)
            target_dir = os.path.join(dst_root, rel_path)
            os.makedirs(target_dir, exist_ok=True)

    def find_docx_files(self, root_folder: str, exclude_prefix: str | None = None) -> List[str]:
        """Поиск всех файлов .docx."""
        if exclude_prefix is None:
            exclude_prefix = DEFAULT_CONFIG.exclude_prefix

        docx_files = []
        for dirpath, _, filenames in os.walk(root_folder):
            for filename in filenames:
                if filename.endswith(".docx") and not filename.startswith(exclude_prefix):
                    filepath = os.path.join(dirpath, filename)
                    docx_files.append(filepath)
        return sorted(docx_files)

    def get_output_path(self, input_path: str, src_root: str, dst_root: str) -> str:
        """Путь для выходного файла."""
        rel_path = os.path.relpath(input_path, src_root)
        return os.path.join(dst_root, rel_path)

    def get_full_text(self, doc: DocumentType) -> str:
        """Получение полного текста."""
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for tc in doc._element.body.iter(qn('w:tc')):
            cell_text = _extract_cell_text(tc)
            if cell_text.strip():
                text_parts.append(cell_text)

        return "\n".join(text_parts)
