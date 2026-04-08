"""
Тесты для модуля docx_processor.
"""

import unittest
import tempfile
import shutil
import os
from docx import Document
from docx.oxml import OxmlElement

from src.date_replacer import DateReplacer
from src.docx_processor import DocxProcessor


class TestDocxProcessor(unittest.TestCase):
    """Тесты для класса DocxProcessor."""

    def setUp(self):
        """Настройка перед каждым тестом."""
        self.replacer = DateReplacer("«29» января 2026 г.", "«26» февраля 2026 г.")
        self.processor = DocxProcessor(self.replacer)
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Очистка после каждого теста."""
        shutil.rmtree(self.test_dir, ignore_errors=True)

    def _create_test_docx(self, content: str, filename: str = "test.docx") -> str:
        """Создаёт тестовый документ .docx."""
        filepath = os.path.join(self.test_dir, filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)
        return filepath

    def test_load_document(self):
        """Загрузка документа."""
        filepath = self._create_test_docx("Тестовый текст")
        doc = self.processor.load_document(filepath)
        self.assertIsNotNone(doc)
        self.assertEqual(len(doc.paragraphs), 1)

    def test_load_document_not_found(self):
        """Загрузка несуществующего документа."""
        with self.assertRaises(FileNotFoundError):
            self.processor.load_document("nonexistent.docx")

    def test_save_document(self):
        """Сохранение документа."""
        doc = Document()
        doc.add_paragraph("Тест")
        filepath = os.path.join(self.test_dir, "output.docx")
        self.processor.save_document(doc, filepath)
        self.assertTrue(os.path.exists(filepath))

    def test_process_paragraphs_with_date(self):
        """Обработка параграфов с датой."""
        doc = Document()
        doc.add_paragraph("«29» января 2026 г.")
        processed, replaced = self.processor.process_paragraphs(doc)
        self.assertGreater(processed, 0)
        self.assertEqual(replaced, 1)
        self.assertIn("26", doc.paragraphs[0].text)
        self.assertIn("февраля", doc.paragraphs[0].text)

    def test_process_paragraphs_without_date(self):
        """Обработка параграфов без даты."""
        doc = Document()
        doc.add_paragraph("«27» февраля 2026 г.")
        processed, replaced = self.processor.process_paragraphs(doc)
        self.assertGreater(processed, 0)
        self.assertEqual(replaced, 0)

    def test_process_paragraphs_split_run(self):
        """Обработка даты, разбитой между несколькими run-ами."""
        doc = Document()
        p = doc.add_paragraph()

        # Run 1: "«29» январ"
        r1 = OxmlElement('w:r')
        t1 = OxmlElement('w:t')
        t1.text = "«29» январ"
        r1.append(t1)
        p._element.append(r1)

        # Run 2: "я 2026 г."
        r2 = OxmlElement('w:r')
        t2 = OxmlElement('w:t')
        t2.text = "я 2026 г."
        r2.append(t2)
        p._element.append(r2)

        processed, replaced = self.processor.process_paragraphs(doc)
        self.assertEqual(replaced, 1)

        full_text = ''.join(run.text for run in doc.paragraphs[0].runs)
        self.assertIn("26", full_text)
        self.assertIn("февраля", full_text)

    def test_process_document_success(self):
        """Успешная обработка документа."""
        input_path = self._create_test_docx("УТВЕРЖДАЮ\n«29» января 2026 г.")
        output_path = os.path.join(self.test_dir, "output.docx")
        success, message, count = self.processor.process_document(input_path, output_path)
        self.assertTrue(success)
        self.assertGreater(count, 0)
        self.assertTrue(os.path.exists(output_path))

    def test_process_document_no_date(self):
        """Обработка документа без даты."""
        input_path = self._create_test_docx("«27» февраля 2026 г.")
        output_path = os.path.join(self.test_dir, "output.docx")
        success, message, count = self.processor.process_document(input_path, output_path)
        self.assertTrue(success)
        self.assertEqual(count, 0)

    def test_find_docx_files(self):
        """Поиск файлов .docx."""
        self._create_test_docx("Тест 1", "file1.docx")
        self._create_test_docx("Тест 2", "file2.docx")
        self._create_test_docx("Тест 3", "~$temp.docx")

        files = self.processor.find_docx_files(self.test_dir)
        self.assertEqual(len(files), 2)
        self.assertNotIn("~$temp.docx", [os.path.basename(f) for f in files])

    def test_get_full_text(self):
        """Получение полного текста документа."""
        doc = Document()
        doc.add_paragraph("Первый параграф")
        doc.add_paragraph("Второй параграф")

        text = self.processor.get_full_text(doc)
        self.assertIn("Первый параграф", text)
        self.assertIn("Второй параграф", text)

    def test_get_full_text_with_tables(self):
        """Получение полного текста: текст из таблиц."""
        doc = Document()
        doc.add_paragraph("Параграф")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Текст из таблицы"

        text = self.processor.get_full_text(doc)
        self.assertIn("Параграф", text)
        self.assertIn("Текст из таблицы", text)

    def test_get_output_path(self):
        """Вычисление выходного пути."""
        src = os.path.join(self.test_dir, "src", "sub")
        dst = os.path.join(self.test_dir, "dst", "sub")
        input_file = os.path.join(src, "file.docx")

        result = self.processor.get_output_path(input_file, src, dst)
        self.assertEqual(result, os.path.join(dst, "file.docx"))


class TestDocxProcessorTables(unittest.TestCase):
    """Тесты для обработки таблиц."""

    def setUp(self):
        """Настройка перед каждым тестом."""
        self.replacer = DateReplacer("«29» января 2026 г.", "«26» февраля 2026 г.")
        self.processor = DocxProcessor(self.replacer)
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Очистка после каждого теста."""
        shutil.rmtree(self.test_dir, ignore_errors=True)

    def test_process_tables_with_date(self):
        """Обработка таблиц с датой."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "УТВЕРЖДАЮ\n«29» января 2026 г."

        filepath = os.path.join(self.test_dir, "table_test.docx")
        doc.save(filepath)

        doc = Document(filepath)
        processed, replaced = self.processor.process_tables(doc)
        self.assertIsNotNone((processed, replaced))

    def test_process_tables_without_date(self):
        """Обработка таблиц без даты."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "«27» февраля 2026 г."

        filepath = os.path.join(self.test_dir, "table_test.docx")
        doc.save(filepath)

        doc = Document(filepath)
        processed, replaced = self.processor.process_tables(doc)
        self.assertIsNotNone((processed, replaced))


if __name__ == "__main__":
    unittest.main()
