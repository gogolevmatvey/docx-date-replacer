"""
Интеграционные тесты для проекта Docx Date Replacer.
"""

import unittest
import tempfile
import shutil
import os
from docx import Document

from src.date_replacer import DateReplacer
from src.docx_processor import DocxProcessor


class TestIntegration(unittest.TestCase):
    """Интеграционные тесты."""

    OLD_DATE = "«29» января 2026 г."
    NEW_DATE = "«26» февраля 2026 г."

    def setUp(self):
        """Настройка перед каждым тестом."""
        self.replacer = DateReplacer(self.OLD_DATE, self.NEW_DATE)
        self.processor = DocxProcessor(self.replacer)
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Очистка после каждого теста."""
        shutil.rmtree(self.test_dir, ignore_errors=True)

    def test_full_workflow_paragraph(self):
        """Полный цикл обработки: параграф с датой."""
        doc = Document()
        doc.add_paragraph("УТВЕРЖДАЮ:")
        doc.add_paragraph("Директор института")
        doc.add_paragraph("__________Паньшин И.А.")
        doc.add_paragraph("«29» января 2026 г.")

        input_path = os.path.join(self.test_dir, "input.docx")
        doc.save(input_path)

        output_path = os.path.join(self.test_dir, "output.docx")
        success, message, count = self.processor.process_document(input_path, output_path)

        self.assertTrue(success)
        self.assertEqual(count, 1)

        result_doc = Document(output_path)
        found_new_date = any("26" in p.text and "февраля" in p.text for p in result_doc.paragraphs)
        self.assertTrue(found_new_date)

    def test_full_workflow_table(self):
        """Полный цикл обработки: таблица с датой."""
        doc = Document()
        table = doc.add_table(rows=4, cols=1)

        table.cell(0, 0).text = "УТВЕРЖДАЮ:"
        table.cell(1, 0).text = "Директор института"
        table.cell(2, 0).text = "__________Паньшин И.А."
        table.cell(3, 0).text = "«29» января 2026 г."

        input_path = os.path.join(self.test_dir, "input.docx")
        doc.save(input_path)

        output_path = os.path.join(self.test_dir, "output.docx")
        success, message, count = self.processor.process_document(input_path, output_path)

        self.assertTrue(success)
        self.assertGreaterEqual(count, 1)

    def test_full_workflow_no_date(self):
        """Полный цикл обработки: документ без даты."""
        doc = Document()
        doc.add_paragraph("УТВЕРЖДАЮ:")
        doc.add_paragraph("«27» февраля 2026 г.")

        input_path = os.path.join(self.test_dir, "input.docx")
        doc.save(input_path)

        output_path = os.path.join(self.test_dir, "output.docx")
        success, message, count = self.processor.process_document(input_path, output_path)

        self.assertTrue(success)
        self.assertEqual(count, 0)

    def test_full_workflow_multiple_dates(self):
        """Полный цикл обработки: несколько дат в документе."""
        doc = Document()
        doc.add_paragraph("«29» января 2026 г.")
        doc.add_paragraph("«27» февраля 2026 г.")  # Эта не должна измениться
        doc.add_paragraph("«29» января 2026 г.")

        input_path = os.path.join(self.test_dir, "input.docx")
        doc.save(input_path)

        output_path = os.path.join(self.test_dir, "output.docx")
        success, message, count = self.processor.process_document(input_path, output_path)

        self.assertTrue(success)
        self.assertGreaterEqual(count, 1)

    def test_folder_structure_preservation(self):
        """Сохранение структуры папок."""
        subdir = os.path.join(self.test_dir, "subdir")
        os.makedirs(subdir)

        doc1 = Document()
        doc1.add_paragraph("«29» января 2026 г.")
        doc1.save(os.path.join(self.test_dir, "file1.docx"))

        doc2 = Document()
        doc2.add_paragraph("«29» января 2026 г.")
        doc2.save(os.path.join(subdir, "file2.docx"))

        output_dir = os.path.join(self.test_dir, "output")
        self.processor.copy_folder_structure(self.test_dir, output_dir)

        self.assertTrue(os.path.exists(os.path.join(output_dir, "subdir")))


if __name__ == "__main__":
    unittest.main()
